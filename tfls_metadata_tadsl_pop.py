# -*- coding: utf-8 -*-
"""
分析集 t14_1-1_2.xlsx 初始化（独立脚本，无 GUI）

从 Word 文档「分析集」章节解析小标题与内容，写入 Excel t14_1-1_2.xlsx。
输入方式：两个输入宏参数，用户调用时在下方设置路径，或通过命令行传入。
无 GUI。可直接运行本脚本或作为模块调用 run_analysis_set_init(sap_docx_path, output_xlsx_path)。

支持的格式（兼容不同作者排版）：
- 章节定位：段落文本含「分析集」即视为分析集章节起始。
- 小标题（TEXT 列来源）：① 带项目符号/编号的段落（Word numPr 或段首 ·•▪▸◆ 等）；② 样式为「标题 2」/ Heading 2 的段落；③ 正文中「标题：内容」单行（无项目符号时也识别，如 SHR-1703-301 等）。
- 下一章边界：样式为「标题 1」且文本不含「分析集」，或段落形如一级编号（如 6. 终点指标）；若分析集对应章节号为 4.1.1.，则 4.1.2（如通用规则）不属于分析集章节，视为下一章，不纳入 excel。
- 单段「标题：内容」格式：若小标题段落内包含中文「：」或英文":"，则符号前写入 TEXT 列，符号后及后续段落写入 FOOTNOTE 列（两种冒号功能相同）。
若新 SAP 格式解析异常，可先运行 inspect_analysis_set_docx.py 查看该文档的分析集段落样式与结构，再按需扩展解析规则。
"""
import logging
import os
import re
import shutil
import sys
from datetime import datetime

logger = logging.getLogger(__name__)

# ---------- 输入宏参数：用户调用时设置以下路径（或通过命令行参数传入） ----------
PATH_SAP_DOCX = ""       # 包含「分析集」章节的 SAP 文档（.docx）完整路径
PATH_T14_1_1_2 = ""      # 输出的 t14_1-1_2.xlsx 完整路径


def _setup_log_file():
    """将本模块日志同时输出到本地文件（logs/tfls_metadata_tadsl_pop.log），仅添加一次。"""
    if any(isinstance(h, logging.FileHandler) for h in logger.handlers):
        return
    log_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logs")
    try:
        os.makedirs(log_dir, exist_ok=True)
        log_path = os.path.join(log_dir, "tfls_metadata_tadsl_pop.log")
        fh = logging.FileHandler(log_path, mode="a", encoding="utf-8")
        fh.setLevel(logging.DEBUG)
        fh.setFormatter(logging.Formatter("%(asctime)s [%(levelname)s] %(name)s %(message)s", datefmt="%Y-%m-%d %H:%M:%S"))
        logger.addHandler(fh)
        logger.setLevel(logging.DEBUG)
    except Exception as e:
        logger.debug("无法创建日志文件 %s: %s", log_dir, e)


_setup_log_file()


# ---------- 分析集 Word 解析与 Excel 生成 ----------
# 支持多种作者排版：章节标题含「分析集」；小标题可为项目符号/编号或「标题 2」样式；下一章为标题 1 或一级编号（如 6. 终点）。

_BULLET_CHARS = (
    "\u2022", "\u2023", "\u25E6", "\u2043", "\u2219", "\u00B7", "\u25AA", "\u25CF",
    "\u25B6", "\u25C6", "\u25A0", "-", "*", "·", "•",
)  # 含 ▪▸◆■ 等常见符号
_BULLET_PATTERN = re.compile(
    r"^[\s\u2022\u2023\u25E6\u2043\u2219\u00B7\u25AA\u25CF\u25B6\u25C6\u25A0\-*\u30FB]+\s*"
)


def _is_bullet_paragraph(paragraph):
    """判断段落是否为项目符号（小黑点）列表项：Word 编号/列表格式或段首为项目符号。"""
    try:
        pPr = paragraph._element.pPr
        if pPr is not None and pPr.numPr is not None:
            return True
    except Exception:
        pass
    text = (paragraph.text or "").strip()
    if not text:
        return False
    return text[0] in _BULLET_CHARS or _BULLET_PATTERN.match(text) is not None


def _strip_bullet(text):
    """去掉段首的小黑点、编号和空白，得到小段标题。支持 (1)、1）、多级编号如 5.1.1。"""
    if not text:
        return text
    t = text.strip()
    t = _BULLET_PATTERN.sub("", t).strip()
    t = re.sub(r"^\(\d+\)\s*", "", t)
    t = re.sub(r"^\d+\)\s*", "", t)
    t = re.sub(r"^\d+(\.\d+)*[\.\)]\s*", "", t)
    t = re.sub(r"^\d+[\.\)\s]+\s*", "", t)
    return t.strip()


def _is_heading2(paragraph):
    """判断段落是否为「标题 2」/ Heading 2 样式（分析集小节可作为小标题）。"""
    try:
        style_name = (paragraph.style and paragraph.style.name) or ""
        s = style_name.strip()
        return bool(re.match(r"^(Heading|标题)\s*2(\s|$|Char)", s, re.I))
    except Exception:
        return False


# 多级章节号（如 4.1.1）：若分析集对应 4.1.1.，则 4.1.2（通用规则等）视为下一章，不纳入 excel
_SECTION_NUM_PATTERN = re.compile(r"^(\d+(?:\.\d+)*)\.?\s*")
_SECTION_NUM_ONLY_PATTERN = re.compile(r"^(\d+(?:\.\d+)*)\.?\s*$")
# 当分析集为 4.1.1 时，段落若以「通用规则」开头（Word 列表编号可能不在 .text 中）也视为下一章
_NEXT_SECTION_4_1_2_TITLE = "通用规则"


def _is_next_chapter_heading(paragraph, section_number=None):
    """
    判断是否为「分析集」之后的下一章节标题。
    仅当为标题 1 且不含「分析集」，或一级编号（如 6. 终点），或同/上级章节号（如 4.1.2 当分析集为 4.1.1）时视为下一章。
    section_number: 从分析集标题提取的章节号（如 "4.1.1"），用于排除 4.1.2 等非分析集小节。
    """
    text = (paragraph.text or "").strip()
    if not text:
        return False
    if "分析集" in text:
        return False
    if text.startswith(_NEXT_SECTION_4_1_2_TITLE):
        return True
    if section_number:
        m = _SECTION_NUM_PATTERN.match(text)
        if m:
            num = m.group(1)
            if num != section_number and not num.startswith(section_number + "."):
                return True
    if re.match(r"^\d+\.\s+(?!\d)", text):
        return True
    try:
        style_name = (paragraph.style and paragraph.style.name) or ""
        s = style_name.strip()
        if re.match(r"^(Heading|标题)\s*1(\s|$|Char)", s, re.I):
            return True
    except Exception:
        pass
    return False


def _is_subtitle_paragraph(paragraph):
    """分析集章节内的小标题：项目符号/编号段落，或「标题 2」样式。"""
    text = (paragraph.text or "").strip()
    if not text:
        return False
    if _is_bullet_paragraph(paragraph):
        return True
    if _is_heading2(paragraph):
        return True
    return False


def parse_analysis_set_from_docx(docx_path):
    """
    从 Word 文档中解析「分析集」章节：按小段标题（项目符号/编号或「标题 2」）、内容拆分为多行。
    返回 list of (小段标题, 内容)，均为字符串。兼容多种作者排版。
    仅解析正文段落（document.paragraphs），不解析表格内内容（特殊表格先不处理，如 SHR-1905-202 等）。
    """
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("请先安装 python-docx：pip install python-docx")

    doc = Document(docx_path)
    paragraphs = doc.paragraphs  # 仅正文段落，不含表格内段落
    start_idx = None
    for i, p in enumerate(paragraphs):
        t = (p.text or "").strip()
        if "分析集" in t:
            start_idx = i
            break
    if start_idx is None:
        raise ValueError("文档中未找到标题内容为「分析集」的章节。")

    heading_text = (paragraphs[start_idx].text or "").strip()
    section_number = None
    m = _SECTION_NUM_PATTERN.match(heading_text)
    if m:
        section_number = m.group(1)
    if section_number is None and start_idx > 0:
        prev_text = (paragraphs[start_idx - 1].text or "").strip()
        m_prev = _SECTION_NUM_ONLY_PATTERN.match(prev_text)
        if m_prev:
            section_number = m_prev.group(1)

    result = []
    i = start_idx + 1
    while i < len(paragraphs):
        p = paragraphs[i]
        text = (p.text or "").strip()

        if _is_next_chapter_heading(p, section_number):
            break

        is_subtitle_like = _is_subtitle_paragraph(p) and text
        pos_cn = text.find("\uFF1A") if "\uFF1A" in text else len(text)
        pos_en = text.find(":") if ":" in text else len(text)
        pos = min(pos_cn, pos_en)
        has_colon = pos < len(text)
        if has_colon:
            before = text[:pos].strip()
            after = text[pos + 1:].strip()
            sub_title = _strip_bullet(before) if is_subtitle_like else before.strip()
            content_parts = [after] if after else []
        else:
            before = ""
            sub_title = _strip_bullet(text) if is_subtitle_like else ""
            content_parts = []

        before_ok = len(before) < 120 if has_colon else True
        if (is_subtitle_like or (text and has_colon and before_ok)) and sub_title:
            j = i + 1
            while j < len(paragraphs):
                q = paragraphs[j]
                q_text = (q.text or "").strip()
                if _is_next_chapter_heading(q, section_number):
                    break
                if _is_subtitle_paragraph(q) and q_text:
                    break
                if q_text and ("\uFF1A" in q_text or ":" in q_text):
                    q_pos_cn = q_text.find("\uFF1A") if "\uFF1A" in q_text else len(q_text)
                    q_pos_en = q_text.find(":") if ":" in q_text else len(q_text)
                    q_pos = min(q_pos_cn, q_pos_en)
                    if q_pos < len(q_text) and len(q_text[:q_pos].strip()) < 120:
                        break
                if q_text:
                    content_parts.append(q_text)
                j += 1
            content = "\n".join(content_parts) if content_parts else ""
            result.append((sub_title, content))
            i = j
        else:
            i += 1

    return result


def _strip_parens(s):
    """去掉字符串中括号及其中的内容，支持中英文括号（）。"""
    if not s:
        return s
    return re.sub(r"[（(].*?[）)]", "", s).strip()


def _backup_existing_to_archive(file_path):
    """
    若 file_path 存在，则复制到同目录下的 99_archive 文件夹，文件名加年月日时分秒后缀。
    返回备份后的路径，若原文件不存在则返回 None。
    """
    if not file_path or not os.path.isfile(file_path):
        return None
    dir_name = os.path.dirname(file_path)
    base_name = os.path.basename(file_path)
    name, ext = os.path.splitext(base_name)
    suffix = datetime.now().strftime("%Y%m%d%H%M%S")
    archive_dir = os.path.join(dir_name, "99_archive")
    os.makedirs(archive_dir, exist_ok=True)
    backup_name = "%s_%s%s" % (name, suffix, ext)
    backup_path = os.path.join(archive_dir, backup_name)
    shutil.copy2(file_path, backup_path)
    return backup_path


def write_tadsl_pop_xlsx(xlsx_path, rows):
    """
    将 (小段标题, 内容) 列表写入 Excel，与完整表格一致：TEXT、ROW、MASK、LINE_BREAK、INDENT、FILTER、FOOTNOTE。
    rows: list of (小段标题, 内容)
    """
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "分析集"
    ws.append(["TEXT", "ROW", "MASK", "LINE_BREAK", "INDENT", "FILTER", "FOOTNOTE"])
    for row_num, (title, content) in enumerate(rows, start=1):
        text_cell = _strip_parens(title)
        footnote = "%s：%s" % (title, content) if content else "%s：" % title
        ws.append([text_cell, row_num, "", "", "", "", footnote])
    d = os.path.dirname(xlsx_path)
    if d:
        os.makedirs(d, exist_ok=True)
    wb.save(xlsx_path)


def run_analysis_set_init(sap_docx_path, output_xlsx_path, backup=True):
    """
    从 SAP 文档解析「分析集」章节并生成 t14_1-1_2.xlsx。

    :param sap_docx_path: 包含「分析集」章节的 SAP 文档（.docx）路径
    :param output_xlsx_path: 输出的 t14_1-1_2.xlsx 路径
    :param backup: 若输出文件已存在，是否先备份到 99_archive
    :return: (success: bool, message: str) 或生成条数（成功时）
    """
    if not sap_docx_path or not sap_docx_path.strip():
        return False, "请设置 PATH_SAP_DOCX（SAP 文档路径）。"
    sap_docx_path = sap_docx_path.strip()
    if not os.path.isfile(sap_docx_path):
        return False, "SAP 文件不存在：%s" % sap_docx_path
    if not output_xlsx_path or not output_xlsx_path.strip():
        return False, "请设置 PATH_T14_1_1_2（t14_1-1_2.xlsx 输出路径）。"
    output_xlsx_path = output_xlsx_path.strip()

    try:
        if backup and os.path.isfile(output_xlsx_path):
            backup_path = _backup_existing_to_archive(output_xlsx_path)
            if backup_path:
                logger.info("已备份原文件至：%s", backup_path)
        rows = parse_analysis_set_from_docx(sap_docx_path)
        if not rows:
            return False, "未在「分析集」章节下解析到任何小段标题与内容。请确认文档中该章节内的小段标题为「小黑点」列表项（项目符号）。"
        write_tadsl_pop_xlsx(output_xlsx_path, rows)
        logger.info("已初始化 t14_1-1_2.xlsx：%s（共 %d 条）", output_xlsx_path, len(rows))
        return True, len(rows)
    except Exception as e:
        logger.exception("分析集初始化失败")
        return False, "解析或生成失败：%s" % e


if __name__ == "__main__":
    # 可选：取消下面注释并改为你的路径后使用宏参数运行
    # PATH_SAP_DOCX = r"Z:\projects\xxx\utility\documentation\03_statistics\SAP.docx"
    # PATH_T14_1_1_2 = r"Z:\projects\xxx\utility\metadata\t14_1-1_2.xlsx"

    #PATH_SAP_DOCX = r".\HRS-2129-101 统计分析计划(SAP) V1.0.docx"
    #PATH_T14_1_1_2 = r".\t14_1-1_2-2129101.xlsx"
    #PATH_SAP_DOCX = r".\SHR-1905-202 统计分析计划-V1.1.docx"
    #PATH_T14_1_1_2 = r".\t14_1-1_2-1905202.xlsx"
    PATH_SAP_DOCX = r"SHR-1703-301SAP V0.3-202601.docx"
    PATH_T14_1_1_2 = r".\t14_1-1_2.xlsx"
    #PATH_SAP_DOCX = r".\HRS-9231-302 统计分析计划.docx"
    #PATH_T14_1_1_2 = r".\t14_1-1_2-9231302.xlsx"

    

    if len(sys.argv) >= 3:
        path_sap = sys.argv[1]
        path_xlsx = sys.argv[2]
    else:
        path_sap = PATH_SAP_DOCX
        path_xlsx = PATH_T14_1_1_2
    if not path_sap:
        print("请设置脚本顶部宏参数 PATH_SAP_DOCX、PATH_T14_1_1_2，或使用：")
        print("  python tfls_metadata_tadsl_pop.py <SAP文档路径> <t14_1-1_2.xlsx输出路径>")
        sys.exit(1)
    if not path_xlsx:
        print("请设置 PATH_T14_1_1_2（t14_1-1_2.xlsx 输出路径）。")
        sys.exit(1)
    ok, result = run_analysis_set_init(path_sap, path_xlsx, backup=True)
    if ok:
        print("分析集 t14_1-1_2.xlsx 已生成，共 %s 条，文件路径为：" % result, path_xlsx)
    else:
        print("失败：%s" % result, file=sys.stderr)
        sys.exit(1)
